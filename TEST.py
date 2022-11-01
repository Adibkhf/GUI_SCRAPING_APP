from PyQt5 import QtCore, QtGui, QtWidgets
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import Liste_des_agences as l
import Trouver_lien as t_l
import pandas as pd 

class WorkerThread_Ecrire(QtCore.QThread):
    def __init__(self,Emplacement,Nom_document,dictio_agence,parent=None):
        super(WorkerThread_Ecrire, self).__init__(parent)
        self.dictio_agence = dictio_agence
        self.d = Document()
        self.d_ar = Document()
        self.Emplacement = Emplacement
        self.Nom_document = Nom_document
    def nouveau_titre(self,doc, titre):  
        H = doc.add_heading(titre, 2)
        title_style = H.style
        H.style.font.color.rgb = RGBColor(31,56, 100)
        title_style.font.size = Pt(36)
        H.paragraph_format.space_before = Pt(250)
        H.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rFonts = title_style.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), "Garamond")
        doc.add_page_break()   
        
    def run(self):         
        dict_funct_ecrire = {"Algérie Présse Service":l.Algerie_press_service,
                             "EFE": l.EFE,
                             "ANADOLU":l.Agencia_Anadolu,
                             "REUTERS":l.REUTERS,
                             "ANSA":l.ANSA,
                             "TASS" : l.TASS,
                             "AFRICA_PRESSE":l.AFRICA_PRESSE}
        for x in self.dictio_agence:
            if len(self.dictio_agence[x])>0:
                self.nouveau_titre(self.d,x)
                for y in self.dictio_agence[x]:
                    try:
                        dict_funct_ecrire[x](y, self.d)
                    except:
                        pass
                p = self.d.add_paragraph('')
                p.paragraph_format.space_after = Pt(3)
                #print("APS 100%")
        if len(self.dictio_agence["AFRICA_PRESSE_AR"])>0:
            self.nouveau_titre(self.d_ar,x)
            for y in self.dictio_agence["AFRICA_PRESSE_AR"]:
                try:
                    l.AFRICA_PRESSE_AR(y, self.d_ar)
                except:
                    pass
            p = self.d_ar.add_paragraph('')
            p.paragraph_format.space_after = Pt(3)
            
        try:  
            self.d.save(self.Emplacement+'/'+self.Nom_document+'.docx')
            self.d_ar.save(self.Emplacement+'/'+self.Nom_document+'_ar.docx')
        except:
            pass
                #print("Agencia Anadolu 100%")    
        #print("fichier excel éxporté")
    
class WorkerThread_Trouver_liens(QtCore.QThread):
    def __init__(self,la_date,checked_agence,checked_mot_clé,chemin,parent=None):
        super(WorkerThread_Trouver_liens, self).__init__(parent)
        self.la_date = la_date
        self.checked_agence =checked_agence
        self.checked_mot_clé = checked_mot_clé
        self.chemin= chemin
        
    def run(self):         
        la_date = self.la_date.getDate()
        res_f = []
        jour = la_date[2]
        mois = la_date[1]
        Année = la_date[0]
        dict_funct_agence = {"Algérie Présse Service":t_l.Liens_APS,
                             "EFE":t_l.Liens_EFE,
                             "Agencia Anadolu":t_l.Liens_ANADOLU,
                             "REUTERS":t_l.Liens_REUTERS,
                             "LUSA":t_l.Liens_LUSA,
                             "AFRICA_PRESSE" : t_l.Liens_AFRICA_PRESS}
       
        for x in self.checked_agence:
            res_f = res_f + dict_funct_agence[x](int(jour), mois,Année, self.checked_mot_clé)
        df1 = pd.DataFrame(res_f, columns=["Date sortie","Nom de l'agence","Titre de l'article","liens"])
        if len(str(jour)) == 1:
            jour = "0"+str(jour)
        if len(str(mois)) == 1:
            mois = "0"+str(mois)
        df1.to_excel(self.chemin+"/RDP_API_"+str(jour)+"_"+str(mois)+"_"+str(Année)+".xlsx")
        
class CheckableComboBox(QtWidgets.QComboBox):
    def __init__(self, parent=None):
        super(CheckableComboBox, self).__init__(parent)
        self.view().pressed.connect(self.handleItemPressed)
    #    self.view().pressed.connect(self.test_check)
        self._changed = False

        
    def handleItemPressed(self, index):
        item = self.model().itemFromIndex(index)
        if item.checkState() == QtCore.Qt.Checked:
            item.setCheckState(QtCore.Qt.Unchecked)
        else:
            item.setCheckState(QtCore.Qt.Checked)
        if self.itemChecked(0):
            item_tout = self.model().item(0, self.modelColumn())
            for x in range(1,self.count()):
                item = self.model().item(x, self.modelColumn())
                item.setCheckState(item_tout.checkState())
                    
        self._changed = True

    def hidePopup(self):
        if not self._changed:
            super(CheckableComboBox, self).hidePopup()
        self._changed = False

    def itemChecked(self, index):
        item = self.model().item(index, self.modelColumn())
        return item.checkState() == QtCore.Qt.Checked

    def setItemChecked(self, index, checked=True):
        item = self.model().item(index, self.modelColumn())
        if checked:
            item.setCheckState(QtCore.Qt.Checked)
        else:
            item.setCheckState(QtCore.Qt.Unchecked)
    def check_items(self):
        checkedItems = []
        for i in range(self.count()):
            if self.itemChecked(i):
                checkedItems.append(self.model().item(i, 0).text())
        return checkedItems
class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.setFixedSize(848, 483)
        MainWindow.setAutoFillBackground(True)
        MainWindow.setStyleSheet("QFrame{\n"
"border : solid 10 px rgba(0,0,0);\n"
"\n"
"    }\n"
"")
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.tabWidget = QtWidgets.QTabWidget(self.centralwidget)
        self.tabWidget.setGeometry(QtCore.QRect(0, 0, 871, 541))
        self.tabWidget.setStyleSheet("QWidget{\n"
"border : solid 10 px rgba(0,0,0);\n"
"background-color :#eee;\n"
"    }\n"
"\n"
"\n"
"QLabel{\n"
"font-weight:bold;\n"
"font-family:Garamond;\n"
"font-size:24px;\n"
"}\n"
"QPushButton{\n"
"background-color : #001f3f;\n"
"border: 2px solid #000000;\n"
"border-radius:10px;\n"
"color:#fff;\n"
"}\n"
"QPushButton:hover{\n"
"background-color:rgb(0, 0, 116);\n"
"}\n"
"\n"
"QTextEdit{\n"
"border-radius:10px;\n"
"border: 2px solid #000000;\n"
"background-color:rgb(255, 255, 255)\n"
"}\n"
"\n"
"QTextEdit:focus{\n"
" border:2px solid rgb(170, 170, 170);\n"
"}\n"
"QComboBox{\n"
"background : #55d1d3;\n"
"border: 2px solid #000000;\n"
"border-radius:10px;\n"
"background-color:rgb(255, 255, 255)\n"
"\n"
"}\n"
"\n"
"QDateEdit{\n"
"background : #55d1d3;\n"
"border: 2px solid #000000;\n"
"border-radius:10px;\n"
"background-color:rgb(255, 255, 255)\n"
"}")
        self.tabWidget.setObjectName("tabWidget")
        self.tab = QtWidgets.QWidget()
        self.tab.setAutoFillBackground(False)
        self.tab.setObjectName("tab")
        self.label_2 = QtWidgets.QLabel(self.tab)
        self.label_2.setGeometry(QtCore.QRect(20, 40, 141, 31))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(-1)
        font.setBold(True)
        font.setWeight(75)
        self.label_2.setFont(font)
        self.label_2.setObjectName("label_2")
        self.label_4 = QtWidgets.QLabel(self.tab)
        self.label_4.setGeometry(QtCore.QRect(20, 240, 191, 31))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(-1)
        font.setBold(True)
        font.setWeight(75)
        self.label_4.setFont(font)
        self.label_4.setObjectName("label_4")
        self.label_5 = QtWidgets.QLabel(self.tab)
        self.label_5.setGeometry(QtCore.QRect(20, 140, 191, 31))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(-1)
        font.setBold(True)
        font.setWeight(75)
        self.label_5.setFont(font)
        self.label_5.setObjectName("label_5")
        self.pushButton = QtWidgets.QPushButton(self.tab,clicked = lambda : New_document())
        self.pushButton.setGeometry(QtCore.QRect(100, 340, 251, 61))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(14)
        self.pushButton.setFont(font)
        self.pushButton.setObjectName("pushButton")
        self.pushButton_2 = QtWidgets.QPushButton(self.tab,clicked = lambda : Ecrire())
        self.pushButton_2.setGeometry(QtCore.QRect(490, 340, 251, 61))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(14)
        self.pushButton_2.setFont(font)
        self.pushButton_2.setObjectName("pushButton_2")
        #self.pushButton_4 = QtWidgets.QPushButton(self.tab)#clicked = lambda : Sauvegarder())
        #♦self.pushButton_4.setGeometry(QtCore.QRect(560, 340, 251, 61))
        #font = QtGui.QFont()
        #font.setFamily("Garamond")
        #font.setPointSize(14)
        #self.pushButton_4.setFont(font)
        #self.pushButton_4.setObjectName("pushButton_4")
        self.pushButton_5 = QtWidgets.QPushButton(self.tab,clicked = lambda : Save_Path())
        self.pushButton_5.setGeometry(QtCore.QRect(680, 140, 41, 41))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(11)
        self.pushButton_5.setFont(font)
        self.pushButton_5.setStyleSheet("background-color : #FFF;\n"
"color : rgb(0, 0, 0);")
        self.pushButton_5.setObjectName("pushButton_5")
        self.pushButton_6 = QtWidgets.QPushButton(self.tab,clicked = lambda : get_file_name())
        self.pushButton_6.setGeometry(QtCore.QRect(680, 40, 41, 41))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(11)
        self.pushButton_6.setFont(font)
        self.pushButton_6.setStyleSheet("background-color : #FFF;\n"
"color : rgb(0, 0, 0);")
        self.pushButton_6.setObjectName("pushButton_6")
        self.textEdit_2 = QtWidgets.QTextEdit(self.tab)
        self.textEdit_2.setGeometry(QtCore.QRect(180, 240, 491, 41))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(16)
        self.textEdit_2.setFont(font)
        self.textEdit_2.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.textEdit_2.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.textEdit_2.setLineWrapMode(QtWidgets.QTextEdit.NoWrap)
        self.textEdit_2.setObjectName("textEdit_2")
        self.textEdit_3 = QtWidgets.QTextEdit(self.tab)
        self.textEdit_3.setGeometry(QtCore.QRect(180, 140, 491, 41))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(16)
        self.textEdit_3.setFont(font)
        self.textEdit_3.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.textEdit_3.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.textEdit_3.setLineWrapMode(QtWidgets.QTextEdit.NoWrap)
        self.textEdit_3.setObjectName("textEdit_3")
        self.textEdit_8 = QtWidgets.QTextEdit(self.tab)
        self.textEdit_8.setGeometry(QtCore.QRect(180, 40, 491, 41))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(16)
        self.textEdit_8.setFont(font)
        self.textEdit_8.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.textEdit_8.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.textEdit_8.setLineWrapMode(QtWidgets.QTextEdit.NoWrap)
        self.textEdit_8.setObjectName("textEdit_8")
        self.tabWidget.addTab(self.tab, "")
        self.tab_2 = QtWidgets.QWidget()
        self.tab_2.setAutoFillBackground(False)
        self.tab_2.setObjectName("tab_2")
        self.label_3 = QtWidgets.QLabel(self.tab_2)
        self.label_3.setGeometry(QtCore.QRect(10, 40, 101, 31))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(-1)
        font.setBold(True)
        font.setWeight(75)
        self.label_3.setFont(font)
        self.label_3.setObjectName("label_3")
        self.date = QtWidgets.QDateEdit(self.tab_2)
        self.date.setGeometry(QtCore.QRect(220, 30, 491, 41))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(14)
        self.date.setFont(font)
        self.date.setDateTime(QtCore.QDateTime(QtCore.QDate(2022, 3, 1), QtCore.QTime(0, 0, 0)))
        self.date.setObjectName("date")
        self.comboBox = CheckableComboBox(self.tab_2)
        data = ('Tout','Algérie Présse Service','EFE',"Agencia Anadolu","REUTERS","LUSA","AFRICA_PRESSE")
        for index, element in enumerate(data):
            self.comboBox.addItem(element)
            item = self.comboBox.model().item(index, 0)
            item.setCheckState(QtCore.Qt.Unchecked)
        self.comboBox.setGeometry(QtCore.QRect(220, 110, 491, 41))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(14)
        self.comboBox.setFont(font)
        self.comboBox.setObjectName("comboBox")
        #self.comboBox.addItem("")
        #self.comboBox.addItem("")
        self.comboBox_4 = CheckableComboBox(self.tab_2)
        data = ('Tout','Maroc','Sahara')
        for index, element in enumerate(data):
            self.comboBox_4.addItem(element)
            item = self.comboBox_4.model().item(index, 0)
            item.setCheckState(QtCore.Qt.Unchecked)
        self.comboBox_4.setGeometry(QtCore.QRect(220, 190, 491, 41))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(14)
        self.comboBox_4.setFont(font)
        self.comboBox_4.setObjectName("comboBox_4")
        self.pushButton_3 = QtWidgets.QPushButton(self.tab_2,clicked = lambda : Trouver_les_articles())
        self.pushButton_3.setGeometry(QtCore.QRect(290, 340, 251, 61))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(14)
        self.pushButton_3.setFont(font)
        self.pushButton_3.setObjectName("pushButton_3")
        self.label_9 = QtWidgets.QLabel(self.tab_2)
        self.label_9.setGeometry(QtCore.QRect(10, 110, 181, 31))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(-1)
        font.setBold(True)
        font.setWeight(75)
        self.label_9.setFont(font)
        self.label_9.setObjectName("label_9")
        self.label_10 = QtWidgets.QLabel(self.tab_2)
        self.label_10.setGeometry(QtCore.QRect(10, 190, 181, 31))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(-1)
        font.setBold(True)
        font.setWeight(75)
        self.label_10.setFont(font)
        self.label_10.setObjectName("label_10")
        self.label_6 = QtWidgets.QLabel(self.tab_2)
        self.label_6.setGeometry(QtCore.QRect(10, 270, 191, 31))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(-1)
        font.setBold(True)
        font.setWeight(75)
        self.label_6.setFont(font)
        self.label_6.setObjectName("label_6")
        self.textEdit_4 = QtWidgets.QTextEdit(self.tab_2)
        self.textEdit_4.setGeometry(QtCore.QRect(220, 270, 491, 41))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(16)
        self.textEdit_4.setFont(font)
        self.textEdit_4.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.textEdit_4.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.textEdit_4.setLineWrapMode(QtWidgets.QTextEdit.NoWrap)
        self.textEdit_4.setObjectName("textEdit_4")
        self.pushButton_13 = QtWidgets.QPushButton(self.tab_2,clicked = lambda : get_file_Path_liens())
        self.pushButton_13.setGeometry(QtCore.QRect(730, 270, 41, 41))
        font = QtGui.QFont()
        font.setFamily("Garamond")
        font.setPointSize(11)
        self.pushButton_13.setFont(font)
        self.pushButton_13.setStyleSheet("background-color : #FFF;\n"
"color : rgb(0, 0, 0);")
        self.pushButton_13.setObjectName("pushButton_13")
        self.tabWidget.addTab(self.tab_2, "")
        MainWindow.setCentralWidget(self.centralwidget)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        self.tabWidget.setCurrentIndex(1)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        df = pd.DataFrame()
    
        
        def New_document():
            global df 
            df = pd.DataFrame()
            self.textEdit_8.setPlainText("")
            self.textEdit_3.setPlainText("")
            self.textEdit_2.setPlainText("")
            #print("c'est fait") 
            msg = QtWidgets.QMessageBox()
            msg.setWindowTitle("OK")
            msg.setText("Nouveau Document")
            msg.setIcon(QtWidgets.QMessageBox.Information)
            msg.setTextFormat(QtCore.Qt.RichText)
            msg.resize(msg.sizeHint())
            x = msg.exec_() 
            
        def evt_worker_finished_trouve():
            msg = QtWidgets.QMessageBox()
            msg.setWindowTitle("OK")
            msg.setText("Lien Trouvé")
            msg.setIcon(QtWidgets.QMessageBox.Information)
            msg.setTextFormat(QtCore.Qt.RichText)
            msg.resize(msg.sizeHint())
            self.pushButton_3.setEnabled(True)
            x = msg.exec_()
            
            
        def Trouver_les_articles():
            #date_dic = {"janvier":"01","février":"02","mars":"03","avril":"04","mai":"05","juin":"06","juillet":"07","août":"08","septembre":"09","octobre":"10","novembre":"11","décembre":"12"}
            checked_agence = self.comboBox.check_items()
            checked_mot_clé = self.comboBox_4.check_items()
            if 'Tout' in checked_agence:
                checked_agence.remove('Tout')
            if 'Tout' in checked_mot_clé:
                checked_mot_clé.remove('Tout')
            la_date = self.date.date()
            chemin = self.textEdit_4.toPlainText()
            self.worker = WorkerThread_Trouver_liens(la_date, checked_agence, checked_mot_clé,chemin)
            self.worker.start()
            self.worker.finished.connect(evt_worker_finished_trouve)
            self.pushButton_3.setEnabled(False)
            

        def Ecrire():
            global df
            dictio_agence = {"Algérie Présse Service":[],
                             "EFE":[],
                             "ANADOLU":[],
                             "REUTERS":[],
                             "ANSA":[],
                             "TASS":[],
                             "AFRICA_PRESSE":[],
                             "AFRICA_PRESSE_AR":[]
                             }
            for x,y in zip(df["Nom de l'agence"],df["liens"]):
                try:
                    dictio_agence[x].append(y)
                except:
                    pass
            Emplacement = self.textEdit_3.toPlainText()
            Nom_document = self.textEdit_2.toPlainText()
            self.worker1 = WorkerThread_Ecrire(Emplacement,Nom_document,dictio_agence)
            self.worker1.start()
            self.worker1.finished.connect(evt_worker_finished_Ecrire)
            self.pushButton_2.setEnabled(False)
        def evt_worker_finished_Ecrire():
            msg = QtWidgets.QMessageBox()
            msg.setWindowTitle("OK")
            msg.setText("Lien écrits")
            msg.setIcon(QtWidgets.QMessageBox.Information)
            msg.setTextFormat(QtCore.Qt.RichText)
            msg.resize(msg.sizeHint())
            x = msg.exec_()
            self.pushButton_2.setEnabled(True)

                    

        def get_file_name():
            global df
            fname = QtWidgets.QFileDialog.getOpenFileName()
            if fname:
                try:
                    df = pd.read_excel(fname[0])
                except:
                     pass
            self.textEdit_8.setPlainText(fname[0])

        def Save_Path():
            fname = QtWidgets.QFileDialog().getExistingDirectory()
            #print(fname)
            self.textEdit_3.setPlainText(fname)
            
        def get_file_Path_liens():
            fname = QtWidgets.QFileDialog().getExistingDirectory()
            #print(fname)
            self.textEdit_4.setPlainText(fname)
                   
            
    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "RDP API"))
        MainWindow.setWindowIcon(QtGui.QIcon('d.ico'))
        self.label_2.setText(_translate("MainWindow", "Liens"))
        self.label_4.setText(_translate("MainWindow", "Nom "))
        self.label_5.setText(_translate("MainWindow", "Emplacement"))
        self.pushButton.setText(_translate("MainWindow", "Nouveau document"))
        self.pushButton_2.setText(_translate("MainWindow", "Ecrire"))
        #self.pushButton_4.setText(_translate("MainWindow", "Sauvegarder"))
        self.pushButton_5.setText(_translate("MainWindow", "..."))
        self.pushButton_6.setText(_translate("MainWindow", "..."))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab), _translate("MainWindow", "Automatisation"))
        self.label_3.setText(_translate("MainWindow", "Date"))
        #self.comboBox.setItemText(0, _translate("MainWindow", "Algérie presse service"))
        #self.comboBox.setItemText(1, _translate("MainWindow", "EFE"))
        #self.comboBox_4.setItemText(0, _translate("MainWindow", "Maroc"))
        #self.comboBox_4.setItemText(1, _translate("MainWindow", "Sahara"))
        self.pushButton_3.setText(_translate("MainWindow", "Trouver"))
        self.label_9.setText(_translate("MainWindow", "Agence"))
        self.label_10.setText(_translate("MainWindow", "Mots clés "))
        self.label_6.setText(_translate("MainWindow", "Emplacement"))
        self.pushButton_13.setText(_translate("MainWindow", "..."))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab_2), _translate("MainWindow", "Trouver les articles"))

if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())

