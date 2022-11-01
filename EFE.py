from bs4 import BeautifulSoup
import requests
from docx.shared import Inches,Cm
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Pt

def get_data(link):
    r = requests.get(link,stream=True,verify = False)#, proxies=proxies)
    content = r.content
    soup = BeautifulSoup(content,'html.parser')
    return soup

#1-Le titre
def titre(S):
    return S.find('h1',attrs = {'class':"entry-title"}).text.strip()

#S = get_data('https://efe.com/espana/que-pasa-barriada-principe-ceuta-escenario-de-docena-tiroteos/')
#r = S.find('div',attrs = {'class':'featured-image page-header-image-single'}).find('img').get('src')

def date(S):
    esp_dict = {'enero': '1','febrero': '2','marzo': '3','abril': '4','mayo': '5','junio': '6','julio': '7','agosto': '8','septiembre': '9','octubre': '10','noviembre': '11','dicembre': '12'}
    date__ = S.find("time").text.split(" ")
    jour = date__[0]
    jour = str(jour)
    if len(jour)==1:
        jour = "0"+jour
    mois = esp_dict[date__[1]]
    Année = date__[2]
    date_de_publication = jour+'/'+mois+'/'+Année
    return date_de_publication



def check_image(S):
    r = S.find('div',attrs = {'class':'featured-image page-header-image-single'})
    if (r is not None):
        response = requests.get(r.find('img').get('src'),verify=False,stream=True)
        file = open("sample_i.png", "wb")
        file.write(response.content)
        file.close()
        return True
    else: 
        return False


#3 les paragraphes 
def Paragraphes(S):
    i = 0 
    d = {} #dictionnaire des paragraphes indexés 
    for x in S.find("div",attrs={"class":"entry-content"}).findAll('p'):
        d[i]=x.text
        i = i+1
    return d



def ecrire_article(S,doc):
    #Titre
    H = doc.add_heading(titre(S), 1)
    title_style = H.style
    title_style.font.size = Pt(24)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Garamond")
    H.alignment = 3     

    #Date et heure de publication
    d = {"janvier":"01","février":"02","mars":"03","avril":"04","mai":"05","juin":"06","juillet":"07","août":"08","septembre":"09","octobre":"10","novembre":"11","décembre":"12"}
    p = doc.add_paragraph('')
    p.add_run('Date de publication ').bold = True
    
    p.add_run(date(S))
    p.alignment = 3
    p = doc.add_paragraph('')
    p.add_run('Heure de publication ').bold = True
    p.add_run(": N.D")
    p.alignment = 3

    #Image
    try:
        if(check_image(S)):
            doc.add_picture('sample_i.png', width=Cm(16.19), height=Inches(3.594488))
    except:
        doc.add_picture('logo-EFE.png',width=Cm(16.19), height=Inches(3.594488))
    
    #Remplissage des paragraphes
    Dict_parag = Paragraphes(S)
    for x in Dict_parag:
        p = doc.add_paragraph('')
        p.add_run(Dict_parag[x]) 
        p.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....
    doc.add_page_break()

def EFE(url,doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(18)
    S = get_data(url)
    ecrire_article(S,doc)
            
    
#S = get_data("https://efe.com/espana/que-pasa-barriada-principe-ceuta-escenario-de-docena-tiroteos/")
#Para = Paragraphes(S)

#URL = "https://efe.com/espana/que-pasa-barriada-principe-ceuta-escenario-de-docena-tiroteos/"
#URL = "https://efe.com/espana/que-pasa-barriada-principe-ceuta-escenario-de-docena-tiroteos/"
#d = Document()
#EFE(URL, d)

#d.save('C:/Users/HP/Desktop/KBscan/RDP/RDP_API_IG/demo.docx')
