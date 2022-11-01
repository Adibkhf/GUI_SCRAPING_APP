from bs4 import BeautifulSoup
import requests
from docx.shared import Inches,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
from docx import Document
from docx.shared import Pt


def get_data(link):
    r = requests.get(link,stream=True,verify = False)#, proxies=proxies)
    content = r.content
    soup = BeautifulSoup(content,'html.parser')
    return soup

#verifie si l'image existe ou pas 
def check_image(S):
    r = S.find('div',attrs = {'class':"itemImageBlock"})
    if (r is not None):
        response = requests.get("https://www.aps.dz"+r.find_all("img")[0].get("src"),verify=False,stream=True)
        file = open("sample_i.png", "wb")
        file.write(response.content)
        file.close()
        return True
    else: 
        return False

#1-Le titre
def titre(S):
    if S.find('div',attrs = {'class':"itemHeader"}).text is not None:
        return S.find('div',attrs = {'class':"itemHeader"}).text.strip()


#2-Date et heure de publication
def date_heure(S):
    date_dic = {"janvier":"01","février":"02","mars":"03","avril":"04","mai":"05","juin":"06","juillet":"07","août":"08","septembre":"09","octobre":"10","novembre":"11","décembre":"12"}
    date_heure = S.find('span',attrs = {'class':'itemDateCreated'}).text
    text_date = date_heure.split()
    jour = text_date[4]
    m = text_date[5]
    mois = date_dic[m]
    Année = text_date[6]
    date_de_publication = jour+'/'+mois+'/'+Année
    heure_de_publication = text_date[7]
    return date_de_publication,heure_de_publication

#3-Texte d'intro en gras
def intro_gras(S):
    r = S.find('div',attrs = {'class':"itemIntroText"})
    if r != None:
        return r.text.replace("\n","")
    else:
        return 0


#4-Liste des paragraphes gras et italic 
def Paragraphes(S):
    i = 0 
    d = {} #dictionnaire des paragraphes indexés 
    em = [] #texte en italic
    strong = [] #texte en gras 
    for x in S.find("div",attrs={'class':'itemFullText'}).findAll("p"):
        if x.find("em") != None :
            em.append(i)
        if x.find("strong") != None :
            strong.append(i)
        d[i]=x.text
        i = i+1
    return d,strong,em

#5-Filtrer les paragraphes
def filtre_dictio_paragraphe(d):
    for x in range(len(d)):
        if d[x] == '\xa0':
            del d[x]
        else:
            d[x] = d[x].replace("\\","")

# -------------------------Ecrire un article à partir d'un site web----------------------
def ecrire_article(S,doc):
    #Titre
    H = doc.add_heading(titre(S), 1)
    title_style = H.style
    title_style.font.size = Pt(24)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Garamond")
    H.alignment = 3     
    
    #Date et heure de publication
    Date_p,Heure_p = date_heure(S)
    p = doc.add_paragraph('')
    p.add_run('Date de publication ').bold = True
    p.add_run(": "+Date_p)
    p.alignment = 3
    p = doc.add_paragraph('')
    p.add_run('Heure de publication ').bold = True
    p.add_run(": "+Heure_p)
    p.alignment = 3

    #Image
    if(check_image(S)):
        doc.add_picture('sample_i.png', width=Cm(16.19), height=Inches(3.594488))
    else:
        doc.add_picture('logo-APS.png')
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #Texte intro en gras
    if intro_gras(S)!=0:
        Text_intro = intro_gras(S)
        p = doc.add_paragraph('')
        p.add_run(Text_intro).bold = True
        p.alignment = 3
    
    #Remplissage des paragraphes
    Dict_parag,strong,em = Paragraphes(S)
    filtre_dictio_paragraphe(Dict_parag)
    for x in Dict_parag:
        if re.search('^Lire aussi', Dict_parag[x])==None:           
            p = doc.add_paragraph('')
            if x in strong and x in em:
                r = p.add_run(Dict_parag[x])
                r.bold = True
                r.italic = True
            elif x in strong:
                p.add_run(Dict_parag[x]).bold = True
            elif x in em :
                p.add_run(Dict_parag[x]).italic = True
            else:
                p.add_run(Dict_parag[x]) 
            p.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....
    doc.add_page_break()


## --------------- Test d'automatisation ---------------------

def Algerie_press_service(link,doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(18)
    S = get_data(link)
    ecrire_article(S,doc)

