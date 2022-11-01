import requests
from bs4 import BeautifulSoup
from docx.shared import Inches
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_data(link):
    r = requests.get(link,stream=True,verify = False)#, proxies=proxies)
    content = r.content
    soup = BeautifulSoup(content,'html.parser')
    return soup


def titre(S):
    return S.find("h1",attrs = {"class":"news-title"}).text.strip()

def date_heure_pub(S):
    dict_date = {"gennaio":"01","febbraio":"02","marzo":"03","aprile":"04","maggio":"05","giugno":"06","luglio":"07","agosto":"08","settembre":"09","ottobre":"10","novembre":"11","dicembre":"12"}
    d = S.find('time').text
    jour = d[1:3]
    mois = dict_date[d.split(" ")[1]]
    année = d.split(" ")[2][:4]
    heure_de_publication = d.split(" ")[2][4:9]
    date_de_publication = jour+'/'+mois+'/'+année
    return date_de_publication,heure_de_publication


def check_image(S):
    r = S.find('div',attrs = {'class':"img-photo ico-60x60"})
    if (r is not None):
        response = requests.get("https://www.ansa.it"+r.find_all("img")[0].get("src"),verify=False,stream=True)
        file = open("Ansa_press.png", "wb")
        file.write(response.content)
        file.close()
        return True
    else: 
        return False

    
# à refaire
def Paragraphes(S):
    for x in S.find('div',attrs = {"class":"span6 pull-right content-news"}).findAll("p"):
        return x.text
    
def ecrire_article(S,doc):
    #Titre
    H = doc.add_heading(titre(S), 1)
    title_style = H.style
    title_style.font.size = Pt(24)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Garamond")
    H.alignment = 3     

    #Date et heure de publication
    Date_p,Heure_p= date_heure_pub(S)
    p = doc.add_paragraph('')
    p.add_run('Date de publication ').bold = True
    p.add_run(": "+Date_p)
    p.alignment = 3
    p = doc.add_paragraph('')
    p.add_run('Heure de publication ').bold = True
    p.add_run(": "+Heure_p)
    p.alignment = 3

    #Image
    if(check_image(S)):
        doc.add_picture('Ansa_press.png', width=Inches(5.9488189), height=Inches(3.594488))
    else:
        doc.add_picture('logo-ANSA.png')
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #Remplissage des paragraphes
    p = doc.add_paragraph('')
    l = Paragraphes(S)
    l = l.replace("\xa0","")
    l = l.replace("\t","")
    l = l.replace("\n"," ")
    p.add_run(l.strip()) 
    p.alignment = 3# for left, 1 for center, 2 right, 3 justify ....
    doc.add_page_break()

#ecrire_article(["https://www.ansa.it/sito/notizie/cronaca/2022/02/27/bracciante-morto-di-freddo-a-barletta-una-targa-lo-ricorda_6101a7b0-0224-4e0e-8c17-1574c99f2095.html"])

def ANSA(link,doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(18)
    S = get_data(link)
    ecrire_article(S,doc)
