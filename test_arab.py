# -*- coding: utf-8 -*-
"""
Created on Mon Mar 28 08:24:59 2022

@author: HP
"""

from docx import Document
from docx.shared import Pt

#path to doc with altered style:
base_doc_location = Document()
doc = Document(base_doc_location)
my_style = doc.styles['Normal']

# define your desired fonts
user_cs_font_size = 16
user_cs_font_name = 'FreeFarsi'
user_en_font_size = 12
user_en_font_name = 'FreeMono'

# get <w:rPr> element of this style
rpr = my_style.element.rPr

#==================================================
'''This probably isn't necessary if you already
have a document with altered style, but just to be
safe I'm going to add this here'''

if rpr.rFonts is None:
    rpr._add_rFonts()
if rpr.sz is None:
    rpr._add_sz()
#==================================================

'''Get the nsmap string for rpr. This is that "w:"
at the start of elements and element values in xml.
Like these:
    <w:rPr>
    <w:rFonts>
    w:val

The nsmap is like a url:
http://schemas.openxmlformats.org/...

Now w:rPr translates to:
{nsmap url string}rPr

So I made the w_nsmap string like this:'''

w_nsmap = '{'+rpr.nsmap['w']+'}'
#==================================================

'''Because I didn't find any better ways to get an
element based on its tag here's a not so great way
of getting it:
'''
szCs = None
lang = None

for element in rpr:
    if element.tag == w_nsmap + 'szCs':
        szCs = element
    elif element.tag == w_nsmap + 'lang':
        lang = element

'''if there is a szCs and lang element in your style
those variables will be assigned to it, and if not
we make those elements and add them to rpr'''

if szCs is None:
    szCs = rpr.makeelement(w_nsmap+'szCs',nsmap=rpr.nsmap)
if lang is None:
    lang = rpr.makeelement(w_nsmap+'lang',nsmap =rpr.nsmap)

rpr.append(szCs)
rpr.append(lang)
#==================================================

'''Now to set our desired values to these elements
we have to get attrib dictionary of these elements
and set the name of value as key and our value as
value for that dict'''

szCs_attrib = szCs.attrib
lang_attrib = lang.attrib
rFonts_atr = rpr.rFonts.attrib

'''sz and szCs values are string values and 2 times
the font size so if you want font size to be 11 you
have to set sz (for western fonts) or szCs (for CTL
fonts) to "22" '''
szCs_attrib[w_nsmap+'val'] =str(int(user_cs_font_size*2))

'''Now to change cs font and bidi lang values'''
rFonts_atr[w_nsmap+'cs'] = user_cs_font_name
lang_attrib[w_nsmap+'bidi'] = 'fa-IR' # For Persian
#==================================================

'''Because we changed default style we don't even
need to set style every time we add a new paragraph
And if you change font name or size the normal way
it won't change these cs values so you can have a
font for CTL language and a different font for
western language
'''
persian_p = doc.add_paragraph('نوشته')
en_font = my_style.font
en_font.name = user_en_font_name
en_font.size = Pt(user_en_font_size)
english_p = doc.add_paragraph('some text')

doc.save('ex.docx')