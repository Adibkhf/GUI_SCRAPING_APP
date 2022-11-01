from bs4 import BeautifulSoup
import requests
from docx.shared import Inches,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
#import json
from docx import Document
from docx.shared import Pt
#from docx.enum.text import WD_COLOR_INDEX

def get_data(link):
    r = requests.get(link,stream=True,verify = False)#, proxies=proxies)
    content = r.content
    soup = BeautifulSoup(content,'html.parser')
    return soup

    
l = 'https://www.reuters.com/business/sustainable-business/next-solar-expansion-drive-needs-hit-higher-potential-markets-maguire-2022-09-28/'
S = get_data(l)
def date_heure(S):
    for x in S.find("time"):
        print(x.text)
    date = [x.text for x in S.find('time')]
    tm = date[1].split(' ')
    dict_date = {"January":"01","February":"02","March":"03","April":"04","May":"05","June":"06","July":"07","August":"08","September":"09","October":"10","November":"11","December":"12"}
    mois = dict_date[tm[0]]
    jour = tm[1][:-1]
    année = tm[2]
    date_de_publication = jour+'/'+mois+'/'+année
    heure_de_publication = date[1].split(" ")[0]
    return date_de_publication,heure_de_publication

def titre(S):
    return S.find('h1').text

def check_image(S):
    r = S.find("img")
    if (r is not None):
        response = requests.get(r.get("src"),verify=False,stream=True)
        file = open("Reuters.png", "wb")
        file.write(response.content)
        file.close()
        return True
    else: 
        return False

def Paragraphes(S):
    i = 0 
    d = {} #dictionnaire des paragraphes indexés 
    
    for x in S.findAll("p"):
        if re.search('^Our Standards:', x.text)==None:
        
            d[i]=x.text
            d[i] = d[i].replace("\\","")
            i = i+1
        else:
            break

    return d

def ecrire_article(S,doc):
  
    H = doc.add_heading(titre(S), 1)
    title_style = H.style
    title_style.font.size = Pt(24)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Garamond")
    H.alignment = 3     

    #Date et heure de publication
    Date_p,Heure_p = date_heure(S)
    p = doc.add_paragraph('')
    p.add_run('Date de publication ').bold = True
    p.add_run(": "+Date_p)
    p.alignment = 3
    p = doc.add_paragraph('')
    p.add_run('Heure de publication ').bold = True
    p.add_run(": "+Heure_p)
    p.alignment = 3
    #Image
    if(check_image(S)):
        doc.add_picture('Reuters.png', width=Cm(16.19), height=Inches(3.594488))
    else :
        doc.add_picture('logo-REUTERS.png')
    #Remplissage des paragraphes
    Dict_parag = Paragraphes(S)
    for x in Dict_parag:
        p = doc.add_paragraph('')
        p.add_run(Dict_parag[x]) 
        p.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....
    doc.add_page_break()
    #doc.save('C:/Users/HP/Desktop/KBscan/RDP/RDP API/ANADOLU/demo.docx')

def REUTERS(link,doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(18)
    S = get_data(link)
    ecrire_article(S,doc)


#d = Document()
#REUTERS("https://www.reuters.com/business/sustainable-business/next-solar-expansion-drive-needs-hit-higher-potential-markets-maguire-2022-09-28/", d)
#d.save('C:/Users/HP/Desktop/KBscan/RDP/RDP_API_Interface_graphique/demo.docx')










