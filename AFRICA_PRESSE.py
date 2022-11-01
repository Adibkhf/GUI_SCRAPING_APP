import requests
from bs4 import BeautifulSoup
from docx.shared import Inches,Cm
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langdetect import detect

def get_data(link):
    r = requests.get(link,stream=True,verify = False)#, proxies=proxies)
    content = r.content
    soup = BeautifulSoup(content,'html.parser')
    return soup

def date_heure_pub(S):
    d = S.find("time",attrs ={'class':'entry-date updated td-module-date'}).text
    if detect(titre(S))!="ar":
        jour = d[0:2]
        mois = d[3:5]
        Année = d[6:10]
        date_publication = jour+'/'+mois+'/'+Année
        heure_publication = d[13:18]
    else:
        jour = d[8:10]
        mois = d[5:7]
        Année = d[0:4]
        date_publication = jour+'/'+mois+'/'+Année
        heure_publication = d[13:18]
    return date_publication,heure_publication
#4-Liste des paragraphes 

def titre(S):
    return S.find("h1",attrs = {"class":"entry-title"}).text.strip()

def check_image(S):
    r = S.find('div',attrs = {'class':"td-post-featured-image"})
    if (r is not None):
        response = requests.get(r.find_all("img")[0].get("src"),verify=False,stream=True)
        file = open("Africa_press.png", "wb")
        file.write(response.content)
        file.close()
        return True
    else: 
        return False
    
    
def Paragraphes(S):
    i = 0 
    d = {} #dictionnaire des paragraphes indexés 
    pref_list = ['[email protected]','For More News And Analysis','يمكنكم متابعة المزيد','Pour plus d’informations et d’analyses']
    for x in S.find('div',attrs = {"class":"td-post-content"}).findAll("p"):
        x = x.text.strip()
        res = [x.startswith(a) for a in pref_list]
        if any(res) == False:
            d[i]=x
            d[i] = d[i].replace("\\","")
            i = i+1
    return d    
    
def ecrire_article(S,doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(18)
    #Titre
    H = doc.add_heading(titre(S), 1)
    title_style = H.style
    title_style.font.size = Pt(24)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Garamond")
    H.alignment = 3     

    #Date et heure de publication
    Date_p,Heure_p= date_heure_pub(S)
    p = doc.add_paragraph('')
    p.add_run('Date de publication ').bold = True
    p.add_run(": "+Date_p)
    p.alignment = 3
    p = doc.add_paragraph('')
    p.add_run('Heure de publication ').bold = True
    p.add_run(": "+Heure_p)
    p.alignment = 3

    #Image
    if(check_image(S)):
        try:
            doc.add_picture('Africa_press.png', width=Cm(16.19), height=Inches(3.594488))
        except:
            pass
    else:
        doc.add_picture('logo-AFRICAPRESS.png')
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    #Remplissage des paragraphes
    Dict_parag = Paragraphes(S)
    for x in Dict_parag:
        p = doc.add_paragraph('')
        p.add_run(Dict_parag[x]) 
        p.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....
    doc.add_page_break()


def AFRICA_PRESSE(link,doc):
    S = get_data(link)
    ecrire_article(S, doc)

#D = Document()
#AFRICA_PRESSE('https://www.africa-press.net/tanzania/all-news/morocco-death-of-5-year-old-boy-trapped-in-well-saddens-entire-country',D)
#D.save('C:/Users/HP/Desktop/KBscan/RDP/RDP_API_Interface_graphique/demo.docx')


#S =get_data('https://www.africa-press.net/tanzania/all-news/morocco-death-of-5-year-old-boy-trapped-in-well-saddens-entire-country')

#d = Paragraphes(S)
#d[8].startswith("For More News And Analysis")

