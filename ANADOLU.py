# -*- coding: utf-8 -*-
"""
Created on Mon Feb 28 09:07:25 2022

@author: HP
"""
import requests
from bs4 import BeautifulSoup
from docx.shared import Inches,Cm
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_data(link):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 6.0; WOW64; rv:24.0) Gecko/20100101 Firefox/24.0',
        'ACCEPT' : 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
        'ACCEPT-ENCODING' : 'gzip, deflate, br',
        'ACCEPT-LANGUAGE' : 'ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7',
        'REFERER' : 'https://www.google.com/'
    }
    r = requests.get(link,stream=True,verify = False,headers = headers)#, proxies=proxies)
    content = r.content
    soup = BeautifulSoup(content,'html.parser')
    return soup


def date_pub(S):
    d = S.find("span",attrs ={'class':'tarih'}).text
    jour = d[0:2]
    mois = d[3:5]
    Année = d[6:10]
    return jour+'/'+mois+'/'+Année
#4-Liste des paragraphes 
def Paragraphes(S):
    i = 0 
    d = {} #dictionnaire des paragraphes indexés 
    for x in S.find('div',attrs = {"class":"detay-icerik"}).findAll("p"):
        d[i]=x.text
        d[i] = d[i].replace("\\","")
        i = i+1
    del d[len(d)-1]
    return d

#1-Le titre
def titre(S):
    return S.find('div',attrs = {'class':"detay-spot-category"}).find("h1").text.strip()

def check_image(S):
    r = S.find('div',attrs = {'class':"col-md-10"})
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 6.0; WOW64; rv:24.0) Gecko/20100101 Firefox/24.0',
        'ACCEPT' : 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
        'ACCEPT-ENCODING' : 'gzip, deflate, br',
        'ACCEPT-LANGUAGE' : 'ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7',
        'REFERER' : 'https://www.google.com/'
    }
    if (r is not None):
        response = requests.get(r.find_all("img")[0].get("src"),verify=False,stream=True,headers = headers)
        file = open("Anadolu.png", "wb")
        file.write(response.content)
        file.close()
        return True
    else: 
        return False


def Agencia_Anadolu(url,doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(18)
    
    S = get_data(url)
    #Titre
    H = doc.add_heading(titre(S), 1)
    title_style = H.style
    title_style.font.size = Pt(24)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Garamond")
    H.alignment = 3     
    
    #Date et heure de publication
    Date_p = date_pub(S)
    p = doc.add_paragraph('')
    p.add_run('Date de publication ').bold = True
    p.add_run(": "+Date_p)
    p.alignment = 3
    p = doc.add_paragraph('')
    p.add_run('Heure de publication ').bold = True
    p.add_run(": N.D")
    p.alignment = 3

    #Image
    if(check_image(S)):
        doc.add_picture('Anadolu.png', width=Cm(16.19), height=Inches(3.594488))
    else:
        doc.add_picture('logo-AA.png')
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #Remplissage des paragraphes
    Dict_parag = Paragraphes(S)
    for x in Dict_parag:
        p = doc.add_paragraph('')
        p.add_run(Dict_parag[x]) 
        p.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....
    doc.add_page_break()

#doc = Document()            
#Agencia_Anadolu('https://www.aa.com.tr/fr/afrique/oci-conseil-des-ministres-des-affaires-%C3%A9trang%C3%A8res-lafrique-%C3%A0-lordre-du-jour/2539428',doc)
#doc.save('C:/Users/HP/Desktop/KBscan/RDP/RDP_API_Interface_graphique/demo.docx')



    






