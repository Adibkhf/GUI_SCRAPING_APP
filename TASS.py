import requests
import datetime
import time
from bs4 import BeautifulSoup
from docx.shared import Inches
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Pt

def get_data(link):
    r = requests.get(link,stream=True,verify = False)#, proxies=proxies)
    content = r.content
    soup = BeautifulSoup(content,'html.parser')
    return soup


def titre(S):
    return S.find("section",attrs = {"id":"news"}).find("h1",attrs= {"class":"news-header__title"}).text.replace("\n","")

def Intro(S):
    return S.find("section",attrs = {"id":"news"}).find("div",attrs= {"class":"news-header__lead"}).text

def Paragraphes(S):
    i = 0 
    d = {}
    for x in S.find("section",attrs = {"id":"news"}).find("div",attrs= {"class":"text-block"}).findAll("p"):
        d[i]=x.text
        i = i+1        
    return d

def date_heure(S):
    t = int(S.find("dateformat").get('time'))
    n_t = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(t))
    année = n_t[0:4]
    mois = n_t[5:7]
    jour = n_t[8:10]
    heure_de_publication = n_t[11:16]
    date_de_publication = jour+'/'+mois+'/'+année
    return date_de_publication,heure_de_publication
    
def check_image(S):
    r = S.find("div",attrs = {'class':'text-include text-include-photo'}).find("img")
    if (r is not None):
        response = requests.get('https:'+r.get("src"),verify=False,stream=True)
        file = open("sample_i.png", "wb")
        file.write(response.content)
        file.close()
        return True
    else: 
        return False

def ecrire_article(S,doc):
    #Titre
    H = doc.add_heading(titre(S), 1)
    title_style = H.style
    title_style.font.size = Pt(24)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Garamond")
    H.alignment = 3     
    
    #Date et heure de publication
    Date_p,Heure_p = date_heure(S)
    p = doc.add_paragraph('')
    p.add_run('Date de publication ').bold = True
    p.add_run(": "+Date_p)
    p.alignment = 3
    p = doc.add_paragraph('')
    p.add_run('Heure de publication ').bold = True
    p.add_run(": "+Heure_p)
    p.alignment = 3

    #Image
    if(check_image(S)):
        doc.add_picture('Anadolu.png', width=Inches(5.9488189), height=Inches(3.594488))
    else :
        doc.add_picture('logo-Russian_news.png')
            
    #Intro
    p = doc.add_paragraph('')
    p.add_run(Intro(S)) 
    p.alignment = 3 # for lef
    
    #Remplissage des paragraphes
    Dict_parag = Paragraphes(S)
    for x in Dict_parag:
        p = doc.add_paragraph('')
        p.add_run(Dict_parag[x]) 
        p.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....
    doc.add_page_break()

def TASS(link,doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(18)
    S = get_data(link)
    ecrire_article(S,doc)

